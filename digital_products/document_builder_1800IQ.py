import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.colors import HexColor
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import BaseDocTemplate, PageTemplate

DISCLAIMER_TEXT = """
**Important Disclaimers:**
• Educational & Motivational Only: This material is strictly for educational and motivational purposes.
• Not Medical or Therapy Advice: This is not therapy, mental-health treatment, or medical advice.
• Not Financial or Legal Advice: This content does not constitute financial or legal advice.
• Responsible AI Use: Any AI guidance provided is for responsible support and brainstorming, not for academic dishonesty, cheating, or misrepresentation.
• Compliance: Student-athletes must strictly follow all school, team, league, NIL, and academic-integrity policies.
"""

INTRO_TEXT = """
Welcome to the Finish Strong series.

As a former elite sprinter, Canadian National Team athlete, and coach, I know exactly what it feels like to stand on the starting line with the weight of the world on your shoulders. I know the sting of setbacks, the grind of daily discipline, and the pressure to perform at a world-class level.

But I also know the power of resilience, faith, and elite mental fortitude.

These materials are designed for you—the elite student-athlete navigating high-stakes environments, million-dollar NIL landscapes, global media scrutiny, and the immense pressure of expectation. 

Over the course of this journey, we will always stay grounded in **The 10 F's**: Faith, Family, Finances, Future, Focus, Fundamentals, Fitness, Fortitude, Friends, and Fun. These pillars are the foundation of your identity, your brand equity, and your purpose beyond the scoreboard. My message to you is simple: no matter how the race begins, you have the power to Finish Strong.

Take time to reflect, prepare, and dominate using the 10 F's framework. You are more than an athlete; you are a global enterprise and a leader.

Finish Strong,
Lornette Daye
"""

OUTRO_TEXT = """
You have completed this portion of the Finish Strong journey!

Remember that the elite lessons you've learned are not just for your sport; they are for your life. The resilience, strategic discipline, and unshakeable faith you are building now will carry you through every major transition, every global challenge, and every victory that lies ahead.

Keep building your identity. Keep protecting your brand. Keep finishing strong.

Contact Information:
Website: https://lornettedaye.com/
"""

LOGO_PATH = r"C:\Users\Zhane\.gemini\antigravity\brain\271dd3a7-e23b-444c-bad0-29631a649534\media__1783067789789.png"

# Colors for Ultimate 1800IQ Design (Zero Black/Charcoal)
SAPPHIRE = HexColor("#0A192F")
SOFT_GOLD = HexColor("#D4AF37")
LIGHT_GOLD = HexColor("#FDFCF7")
ACCENT_GRAY = HexColor("#708090") # Slate gray, no black

def header_footer(canvas, doc):
    canvas.saveState()
    
    # Top Gold Bar
    canvas.setFillColor(SOFT_GOLD)
    canvas.rect(0, letter[1] - 15, letter[0], 15, fill=1, stroke=0)
    
    # Bottom Footer
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(SAPPHIRE)
    canvas.drawString(50, 30, "© 2026 Lornette Daye. All rights reserved. | lornettedaye.com")
    canvas.drawRightString(letter[0] - 50, 30, f"Page {doc.page}")
    
    canvas.restoreState()

class DocumentBuilder:
    def __init__(self, title, output_dir, filename_base):
        self.title = title
        # Prepend the new folder path to avoid overwriting the originals
        self.output_dir = os.path.join(r"C:\Users\Zhane\Documents\New project\10F\1800IQ_Elite_Products", output_dir)
        self.filename_base = filename_base
        
        self.docx_gold = RGBColor(212, 175, 55)
        self.docx_sapphire = RGBColor(10, 25, 47)
        
        self.doc = Document()
        self.pdf_elements = []
        
        self.styles = getSampleStyleSheet()
        
        # 1800IQ Typography
        self.styles.add(ParagraphStyle(name='TitlePremium', parent=self.styles['Heading1'], alignment=TA_CENTER, fontSize=32, leading=38, spaceAfter=25, textColor=SOFT_GOLD, fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='AuthorPremium', parent=self.styles['Normal'], alignment=TA_CENTER, fontSize=14, spaceAfter=50, textColor=SAPPHIRE, fontName='Helvetica-Bold', textTransform='uppercase'))
        self.styles.add(ParagraphStyle(name='NormalPremium', parent=self.styles['Normal'], fontSize=11, spaceAfter=16, leading=18, textColor=SAPPHIRE))
        self.styles.add(ParagraphStyle(name='BoldPremium', parent=self.styles['NormalPremium'], fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='ItalicPremium', parent=self.styles['NormalPremium'], fontName='Helvetica-Oblique', textColor=ACCENT_GRAY))
        
        self.styles.add(ParagraphStyle(name='Heading2Premium', parent=self.styles['Heading2'], fontSize=22, spaceAfter=20, leading=26, textColor=SOFT_GOLD, fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='Heading3Premium', parent=self.styles['Heading3'], fontSize=16, spaceAfter=14, textColor=SAPPHIRE, fontName='Helvetica-Bold'))
        
        self.styles.add(ParagraphStyle(name='FactBoxTitle', parent=self.styles['NormalPremium'], fontSize=10, textColor=SOFT_GOLD, fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='FactBoxBody', parent=self.styles['NormalPremium'], fontSize=11, leading=16, textColor=SAPPHIRE, fontName='Helvetica-Oblique'))

    def _add_docx_paragraph(self, text, style=None, bold=False, italic=False, center=False, is_heading=False, is_subheading=False):
        p = self.doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if is_heading:
            run.font.color.rgb = self.docx_gold
            run.font.size = Pt(18)
            run.bold = True
        elif is_subheading:
            run.font.color.rgb = self.docx_sapphire
            run.font.size = Pt(14)
            run.bold = True
        else:
            run.font.color.rgb = self.docx_sapphire
            run.font.size = Pt(11)
            
        if style:
            p.style = style
        return p

    def add_title_page(self, author="Lornette Daye"):
        # DOCX Title
        self._add_docx_paragraph("\n\n\n")
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(LOGO_PATH, width=Inches(3.5))
            
        self._add_docx_paragraph("\n")
        p = self._add_docx_paragraph(self.title.upper(), center=True)
        p.runs[0].font.size = Pt(28)
        p.runs[0].font.color.rgb = self.docx_gold
        p.runs[0].bold = True
        
        p = self._add_docx_paragraph(f"BY {author.upper()}", center=True)
        p.runs[0].bold = True
        self.doc.add_page_break()
        
        # PDF Title
        self.pdf_elements.append(Spacer(1, 120))
        if os.path.exists(LOGO_PATH):
            self.pdf_elements.append(Image(LOGO_PATH, width=220, height=220))
            self.pdf_elements.append(Spacer(1, 40))
            
        self.pdf_elements.append(Paragraph(self.title.upper(), self.styles['TitlePremium']))
        self.pdf_elements.append(Spacer(1, 20))
        self.pdf_elements.append(Paragraph(f"BY {author.upper()}", self.styles['AuthorPremium']))
        self.pdf_elements.append(PageBreak())

    def add_disclaimer_and_intro(self):
        self.add_heading("Copyright & Disclaimers", level=2)
        copyright_text = [
            "© 2026 Lornette Daye. All rights reserved.",
            "This material is protected by trademark and copyright law. It may not be reproduced, stored in a retrieval system, or transmitted in any form or by any means—electronic, mechanical, photocopy, recording, or otherwise—without prior written permission from the publisher.",
            "",
            "Contact Information:",
            "Website: https://lornettedaye.com/"
        ]
        for line in copyright_text:
            self.add_paragraph(line)
        self.add_paragraph("")
        for line in DISCLAIMER_TEXT.strip().split('\n'):
            self.add_paragraph(line)
            
        self.add_page_break()
        
        self.add_heading("Introduction", level=2)
        for line in INTRO_TEXT.strip().split('\n'):
            if line:
                self.add_paragraph(line)
        self.add_page_break()

    def add_outro(self):
        self.add_heading("Closing Encouragement", level=2)
        for line in OUTRO_TEXT.strip().split('\n'):
            if line:
                self.add_paragraph(line)

    def add_heading(self, text, level=2):
        if level == 2:
            self._add_docx_paragraph(text, is_heading=True)
            self.pdf_elements.append(Paragraph(text, self.styles['Heading2Premium']))
        else:
            self._add_docx_paragraph(text, is_subheading=True)
            self.pdf_elements.append(Paragraph(text, self.styles['Heading3Premium']))

    def add_paragraph(self, text, bold=False, italic=False, center=False):
        if not text.strip():
            self.doc.add_paragraph()
            self.pdf_elements.append(Spacer(1, 10))
            return

        self._add_docx_paragraph(text, bold=bold, italic=italic, center=center)
        
        style_name = 'AuthorPremium' if center else 'BoldPremium' if bold else 'ItalicPremium' if italic else 'NormalPremium'
        
        safe_text = text
        if "**" in safe_text:
            parts = safe_text.split("**")
            for i in range(1, len(parts), 2):
                parts[i] = f"<b>{parts[i]}</b>"
            safe_text = "".join(parts)
            
        self.pdf_elements.append(Paragraph(safe_text, self.styles[style_name]))
        
    def add_fact_box(self, text):
        # DOCX
        p = self.doc.add_paragraph()
        run = p.add_run(f"REAL WORLD FACT: {text}")
        run.italic = True
        run.bold = True
        run.font.color.rgb = self.docx_gold
        
        # PDF - 1800IQ Premium Box
        self.pdf_elements.append(Spacer(1, 15))
        
        content = [
            Paragraph("DATA & INSIGHT", self.styles['FactBoxTitle']),
            Spacer(1, 8),
            Paragraph(text, self.styles['FactBoxBody'])
        ]
        
        table = Table([[content]], colWidths=[450])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), LIGHT_GOLD),
            ('BOX', (0,0), (-1,-1), 1.5, SOFT_GOLD),
            ('TOPPADDING', (0,0), (-1,-1), 15),
            ('BOTTOMPADDING', (0,0), (-1,-1), 15),
            ('LEFTPADDING', (0,0), (-1,-1), 20),
            ('RIGHTPADDING', (0,0), (-1,-1), 20),
        ]))
        self.pdf_elements.append(table)
        self.pdf_elements.append(Spacer(1, 15))

    def add_table(self, headers, rows):
        # DOCX
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = str(item)
        self.doc.add_paragraph()
        
        # PDF - 1800IQ Premium Table
        data = [headers] + rows
        
        formatted_data = []
        for r_idx, r in enumerate(data):
            formatted_row = []
            for cell in r:
                if r_idx == 0:
                    formatted_row.append(Paragraph(f"<font color='white'>{str(cell)}</font>", self.styles['BoldPremium']))
                else:
                    formatted_row.append(Paragraph(str(cell), self.styles['NormalPremium']))
            formatted_data.append(formatted_row)
            
        # Remove explicit colWidths to let ReportLab auto-calculate and prevent squishing
        pdf_table = Table(formatted_data, colWidths=None)
        pdf_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), SOFT_GOLD), # Gold header instead of Dark Charcoal
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('TOPPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.white),
            ('GRID', (0,0), (-1,-1), 0.5, ACCENT_GRAY),
            ('VALIGN',(0,0),(-1,-1),'TOP'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, LIGHT_GOLD]), # Alternating rows
        ]))
        self.pdf_elements.append(pdf_table)
        self.pdf_elements.append(Spacer(1, 15))

    def add_page_break(self):
        self.doc.add_page_break()
        self.pdf_elements.append(PageBreak())
        
    def add_worksheet_space(self, count=2):
        for _ in range(count):
            self.doc.add_paragraph()
            
        # Add a subtle dotted line for writing in PDF
        for _ in range(count):
            self.pdf_elements.append(Spacer(1, 25))
            self.pdf_elements.append(Paragraph("<font color='#CCCCCC'>....................................................................................................................................................</font>", self.styles['NormalPremium']))

    def save(self):
        os.makedirs(self.output_dir, exist_ok=True)
        docx_path = os.path.join(self.output_dir, f"{self.filename_base}.docx")
        pdf_path = os.path.join(self.output_dir, f"{self.filename_base}.pdf")
        
        self.doc.save(docx_path)
        
        # Build 1800IQ PDF with headers/footers
        doc = BaseDocTemplate(pdf_path, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        template = PageTemplate(id='test', frames=frame, onPage=header_footer)
        doc.addPageTemplates([template])
        doc.build(self.pdf_elements)
        
        print(f"Saved: {docx_path} and {pdf_path}")
