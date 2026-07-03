import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.colors import HexColor
from reportlab.lib import colors

DISCLAIMER_TEXT = """
**Important Disclaimers:**
• Educational & Motivational Only: This material is strictly for educational and motivational purposes.
• Not Medical or Therapy Advice: This is not therapy, mental-health treatment, or medical advice.
• Not Financial or Legal Advice: This content does not constitute financial or legal advice.
• Responsible AI Use: Any AI guidance provided is for responsible support and brainstorming, not for academic dishonesty, cheating, or misrepresentation.
• Compliance: Student-athletes must strictly follow all school, team, league, NIL, and academic-integrity policies.
"""

INTRO_TEXT = """
Welcome to the Finish Strong series.

As a former elite sprinter, Canadian National Team athlete, and coach, I know exactly what it feels like to stand on the starting line with the weight of the world on your shoulders. I know the sting of setbacks, the grind of daily discipline, and the pressure to perform at a world-class level.

But I also know the power of resilience, faith, and elite mental fortitude.

These materials are designed for you—the elite student-athlete navigating high-stakes environments, million-dollar NIL landscapes, global media scrutiny, and the immense pressure of expectation. Over the course of this journey, we will focus on building your identity, your brand equity, and your purpose beyond the scoreboard. My message to you is simple: no matter how the race begins, you have the power to Finish Strong.

Take time to reflect, prepare, and dominate. You are more than an athlete; you are a global enterprise and a leader.

Finish Strong,
Lornette Daye
"""

OUTRO_TEXT = """
You have completed this portion of the Finish Strong journey!

Remember that the elite lessons you've learned are not just for your sport; they are for your life. The resilience, strategic discipline, and unshakeable faith you are building now will carry you through every major transition, every global challenge, and every victory that lies ahead.

Keep building your identity. Keep protecting your brand. Keep finishing strong.

Contact Information:
Website: https://lornettedaye.com/
"""

LOGO_PATH = r"C:\Users\Zhane\.gemini\antigravity\brain\271dd3a7-e23b-444c-bad0-29631a649534\media__1783067789789.png"

class DocumentBuilder:
    def __init__(self, title, output_dir, filename_base):
        self.title = title
        self.output_dir = output_dir
        self.filename_base = filename_base
        
        # Colors
        self.charcoal = HexColor("#333333")
        self.soft_gold = HexColor("#D4AF37")
        self.light_gold = HexColor("#F9F6ED")
        self.docx_gold = RGBColor(212, 175, 55)
        self.docx_charcoal = RGBColor(51, 51, 51)
        
        # Docx Setup
        self.doc = Document()
        
        # ReportLab Setup
        self.pdf_elements = []
        self.styles = getSampleStyleSheet()
        self.styles.add(ParagraphStyle(name='CenterTitle', parent=self.styles['Heading1'], alignment=TA_CENTER, fontSize=28, spaceAfter=20, textColor=self.soft_gold, fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='CenterAuthor', parent=self.styles['Normal'], alignment=TA_CENTER, fontSize=16, spaceAfter=40, textColor=self.charcoal))
        self.styles.add(ParagraphStyle(name='NormalText', parent=self.styles['Normal'], fontSize=12, spaceAfter=14, leading=18, textColor=self.charcoal))
        self.styles.add(ParagraphStyle(name='BoldText', parent=self.styles['NormalText'], fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='ItalicText', parent=self.styles['NormalText'], fontName='Helvetica-Oblique'))
        self.styles.add(ParagraphStyle(name='Heading2_Custom', parent=self.styles['Heading2'], fontSize=20, spaceAfter=16, textColor=self.soft_gold))
        self.styles.add(ParagraphStyle(name='Heading3_Custom', parent=self.styles['Heading3'], fontSize=16, spaceAfter=12, textColor=self.charcoal, fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='FactBoxText', parent=self.styles['NormalText'], fontSize=11, leftIndent=20, rightIndent=20, textColor=self.charcoal, fontName='Helvetica-Oblique'))

    def _add_docx_paragraph(self, text, style=None, bold=False, italic=False, center=False, is_heading=False, is_subheading=False):
        p = self.doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if is_heading:
            run.font.color.rgb = self.docx_gold
            run.font.size = Pt(16)
            run.bold = True
        elif is_subheading:
            run.font.color.rgb = self.docx_charcoal
            run.font.size = Pt(14)
            run.bold = True
        else:
            run.font.color.rgb = self.docx_charcoal
            run.font.size = Pt(11)
            
        if style:
            p.style = style
        return p

    def add_title_page(self, author="Lornette Daye"):
        # DOCX
        self._add_docx_paragraph("\n\n")
        if os.path.exists(LOGO_PATH):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(LOGO_PATH, width=Inches(3.0))
            
        self._add_docx_paragraph("\n")
        p = self._add_docx_paragraph(self.title, center=True)
        p.runs[0].font.size = Pt(24)
        p.runs[0].font.color.rgb = self.docx_gold
        self._add_docx_paragraph(f"By {author}", center=True)
        self.doc.add_page_break()
        
        # PDF
        self.pdf_elements.append(Spacer(1, 100))
        if os.path.exists(LOGO_PATH):
            self.pdf_elements.append(Image(LOGO_PATH, width=200, height=200))
            self.pdf_elements.append(Spacer(1, 40))
            
        self.pdf_elements.append(Paragraph(self.title, self.styles['CenterTitle']))
        self.pdf_elements.append(Paragraph(f"By {author}", self.styles['CenterAuthor']))
        self.pdf_elements.append(PageBreak())

    def add_disclaimer_and_intro(self):
        self.add_heading("Copyright & Disclaimers", level=2)
        
        copyright_text = [
            "© 2026 Lornette Daye. All rights reserved.",
            "This material is protected by trademark and copyright law. It may not be reproduced, stored in a retrieval system, or transmitted in any form or by any means—electronic, mechanical, photocopy, recording, or otherwise—without prior written permission from the publisher.",
            "",
            "Contact Information:",
            "Website: https://lornettedaye.com/"
        ]
        for line in copyright_text:
            self.add_paragraph(line)
            
        self.add_paragraph("")
        for line in DISCLAIMER_TEXT.strip().split('\n'):
            self.add_paragraph(line)
            
        self.add_page_break()
        
        self.add_heading("Introduction from Lornette Daye", level=2)
        for line in INTRO_TEXT.strip().split('\n'):
            if line:
                self.add_paragraph(line)
        self.add_page_break()

    def add_outro(self):
        self.add_heading("Closing Encouragement", level=2)
        for line in OUTRO_TEXT.strip().split('\n'):
            if line:
                self.add_paragraph(line)

    def add_heading(self, text, level=2):
        # DOCX
        if level == 2:
            self._add_docx_paragraph(text, is_heading=True)
        else:
            self._add_docx_paragraph(text, is_subheading=True)
        
        # PDF
        style_name = 'Heading1' if level == 1 else 'Heading2_Custom' if level == 2 else 'Heading3_Custom'
        self.pdf_elements.append(Paragraph(text, self.styles[style_name]))

    def add_paragraph(self, text, bold=False, italic=False, center=False):
        if not text.strip():
            self.doc.add_paragraph()
            self.pdf_elements.append(Spacer(1, 10))
            return

        # DOCX
        self._add_docx_paragraph(text, bold=bold, italic=italic, center=center)
        
        # PDF
        if center:
            style_name = 'CenterAuthor'
        else:
            style_name = 'BoldText' if bold else 'ItalicText' if italic else 'NormalText'
        
        # Very basic markdown bold replacement for PDF
        safe_text = text
        if "**" in safe_text:
            parts = safe_text.split("**")
            # This is a very naive replacement, assuming balanced **
            for i in range(1, len(parts), 2):
                parts[i] = f"<b>{parts[i]}</b>"
            safe_text = "".join(parts)
            
        self.pdf_elements.append(Paragraph(safe_text, self.styles[style_name]))
        
    def add_fact_box(self, text):
        # DOCX (Simulate a box with shading or just indent/italic)
        p = self.doc.add_paragraph()
        run = p.add_run(f"REAL WORLD FACT: {text}")
        run.italic = True
        run.bold = True
        run.font.color.rgb = self.docx_gold
        
        # PDF
        self.pdf_elements.append(Spacer(1, 10))
        data = [[Paragraph(f"<b>REAL WORLD FACT:</b><br/>{text}", self.styles['FactBoxText'])]]
        table = Table(data, colWidths=[400])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), self.light_gold),
            ('BOX', (0,0), (-1,-1), 1, self.soft_gold),
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ]))
        self.pdf_elements.append(table)
        self.pdf_elements.append(Spacer(1, 10))

    def add_table(self, headers, rows):
        # DOCX
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = str(item)
        self.doc.add_paragraph()
        
        # PDF
        data = [headers] + rows
        
        # Convert strings to Paragraphs for text wrapping
        formatted_data = []
        for r in data:
            formatted_row = []
            for cell in r:
                formatted_row.append(Paragraph(str(cell), self.styles['NormalText']))
            formatted_data.append(formatted_row)
            
        # Calculate roughly equal column widths
        col_width = 450 / len(headers)
        
        pdf_table = Table(formatted_data, colWidths=[col_width]*len(headers))
        pdf_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), self.light_gold),
            ('TEXTCOLOR', (0,0), (-1,0), self.charcoal),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.white),
            ('GRID', (0,0), (-1,-1), 1, self.soft_gold),
            ('VALIGN',(0,0),(-1,-1),'TOP'),
        ]))
        self.pdf_elements.append(pdf_table)
        self.pdf_elements.append(Spacer(1, 14))

    def add_page_break(self):
        # DOCX
        self.doc.add_page_break()
        # PDF
        self.pdf_elements.append(PageBreak())
        
    def add_worksheet_space(self, count=2):
        # Better than ugly lines.
        for _ in range(count):
            self.doc.add_paragraph()
            self.pdf_elements.append(Spacer(1, 30))

    def save(self):
        os.makedirs(self.output_dir, exist_ok=True)
        docx_path = os.path.join(self.output_dir, f"{self.filename_base}.docx")
        pdf_path = os.path.join(self.output_dir, f"{self.filename_base}.pdf")
        
        self.doc.save(docx_path)
        
        pdf = SimpleDocTemplate(pdf_path, pagesize=letter, topMargin=40, bottomMargin=40, leftMargin=50, rightMargin=50)
        pdf.build(self.pdf_elements)
        print(f"Saved: {docx_path} and {pdf_path}")
